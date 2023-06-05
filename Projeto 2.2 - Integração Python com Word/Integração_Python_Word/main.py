#Importando a Biblioteca WORD e de Formatação
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#Criando Documento
documento = Document()

#Linha para adicionar textos
texto = """ Fala Adriano,
 Beleza? Como estão as coisas por ai?
 Tamo junto, nós!
 """

paragrafo = documento.add_paragraph(texto)

#Formatação do texto
paragrafo.style = documento.styles.add_style("Estilo Inicial", WD_STYLE_TYPE.PARAGRAPH)
#Estilo inicial é criado com os dados de nome, tamanho, negrito e italico alterando o texto.
paragrafo.style.font.name = "Algerian"
paragrafo.style.font.size = Pt(15)
paragrafo.style.font.bold = True
paragrafo.style.font.italic = True
paragrafo.style.font.color.rgb = RGBColor(255, 0, 0)

#Usando o estilo inicial em outro paragrafo.
paragrafo = documento.add_paragraph("A Quantidade vendida foi de 10", "Estilo Inicial")

#Adicionando texto e variavéis dentro de textos
faturamento = 1000
texto = "Ontem na empresa o faturamento foi de "
texto_final = "Resultando assim um superávit nas contas."

paragrafo = documento.add_paragraph(texto)
paragrafo.add_run(f"R${faturamento}").bold = True #Adiciona a variavél no texto é possível formatar
paragrafo = documento.add_paragraph(texto_final)

#Controle de Margem e Seções
for secao in documento.sections:
    secao.top_margin = Cm(0.5)
    secao.bottom_margin = Cm(1)
    secao.left_margin = Cm(1)
    secao.right_margin = Cm(1)

#Formatar posição de texto
paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

imagem = documento.add_picture("imagem.png", width=Cm(10), height=Cm(10))
imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#Adicionando uma Tabela
records = {
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, Spam, Eggs, and Spam')
}

table = documento.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


#Salvar o arquivo WORD com nome Texto
documento.save("Texto.docx")

#Lista os estilos que estão salvos no WORD
#for estilo in documento.styles:
    #print(estilo)

#Usando um template criado no WORD
#Criar um novo arquivo no word em que tenha criado templates de texto. Exemplo com normas ABNT.
#template = Document("template.docx")
#paragrafo = template.add_paragraph("Testando novo template", "NovoTemplate")
#template.save("NovoArquivoComTemplate.dox")

#ALTERANDO UM DOCUMENTO SALVO
from datetime import datetime
contrato = Document("contrato.docx")

#Novos campos para alterar no documento
nome = "Adriano"
item1 = "Serviço de Treinamento Word"
item2 = "Apostila Completa Word"
item3 = "Serviço de Treinamento Python"

#Estrutura do Contrato alterando os campos
dicionario_valores = {
                      "XXXX":nome,
                      "YYYY":item1,
                      "ZZZZ":item2,
                      "WWWW":item3,
                      "DD":str(datetime.now().day),
                      "MM":str(datetime.now().month),
                      "AAAA":str(datetime.now().year)
                      }
for paragrafo in contrato.paragraphs:
    for codigo in dicionario_valores:
        if codigo in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(codigo, dicionario_valores[codigo])

#Salvando em um novo documento
contrato.save(f"Contrato Atualizado - {nome}.docx")